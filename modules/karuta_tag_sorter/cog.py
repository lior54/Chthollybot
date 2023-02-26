import os
import nextcord
from nextcord.ext import commands
import modules.karuta_tag_sorter.sorter
import docx
from docx.shared import RGBColor

class Karuta(commands.Cog, name="Karuta"):
    def __init__(self, bot: commands.Bot):
        self.bot = bot

    @nextcord.slash_command(name="karuta_sorter", description="Recommending for sort based on ksheet, each series to be tagged must have at least one card tagged")
    async def karuta_sorter(self, interaction: nextcord.Interaction, link: str = nextcord.SlashOption(required=True, name="link", description="link to ksheet file"), exclude: str = nextcord.SlashOption(required=False, name="exclude", description="tags to exclude from tag suggestions saperated by ,")):
        if not link.startswith("https://") and link.endswith(".csv"):
            await interaction.response.send_message("invalid link")
        else:
            if exclude:
                exclude = exclude.replace(" ", "")
                if "," in exclude:
                    exclude = exclude.split(",")
                else:
                    exclude = []
                    exclude.append(exclude)
            else:
                exclude = []
            exclude.append("")
            data = modules.karuta_tag_sorter.sorter.karuta_sorter(link=link, exclude=exclude)
            with open(f"{interaction.user.id}.txt", "w", encoding="utf-8") as file:
                file.write(data)
            await interaction.response.send_message(file=nextcord.File(fp=f"{interaction.user.id}.txt"))
            os.remove(f"{interaction.user.id}.txt")

    @nextcord.slash_command(name="karuta_duplicates", description="Returning list of  duplicated cards")
    async def karuta_duplicates(self, interaction: nextcord.Interaction, link: str = nextcord.SlashOption(required=True, name="link", description="link to ksheet file")):
        if not link.startswith("https://") and link.endswith(".csv"):
            await interaction.response.send_message("invalid link")
        else:
            data = modules.karuta_tag_sorter.sorter.karuta_duplicates(link=link)
            doc = docx.Document()
            paragraph = doc.add_paragraph()
            for row in data.splitlines():
                row = row.split(", ", 1)
                run = paragraph.add_run(row[0] + ", ")
                run.font.color.rgb = RGBColor.from_string("0066CC")
                temp = row[1].rsplit(": ", 1)
                series = temp[0]
                run = paragraph.add_run(series + ": ")
                run.font.color.rgb = RGBColor.from_string("FF0000")
                temp = temp[1]
                temp = temp.split(", ")
                if type(temp) == list:
                    first = temp[0]
                else:
                    first = temp
                run = paragraph.add_run(first[:first.index("(")])
                run.font.color.rgb = RGBColor.from_string("00FF00")
                run = paragraph.add_run(first[first.index("("):])
                run.font.color.rgb = RGBColor.from_string("000000") 
                if type(temp) == list:
                    for edition in temp[1:]:
                        run = paragraph.add_run(", " + edition[:edition.index("(")])
                        run.font.color.rgb = RGBColor.from_string("00FF00")
                        run = paragraph.add_run(edition[edition.index("("):])
                        run.font.color.rgb = RGBColor.from_string("000000")
                run.add_break()
            doc.save(f"{interaction.user.id}.docx")
            await interaction.response.send_message(file=nextcord.File(fp=f"{interaction.user.id}.docx"))
            os.remove(f"{interaction.user.id}.docx")

    @nextcord.slash_command(name="karuta_series", description="returning list of series in collection")
    async def karuta_sorter(self, interaction: nextcord.Interaction, link: str = nextcord.SlashOption(required=True, name="link", description="link to ksheet file")):
        if not link.startswith("https://") and link.endswith(".csv"):
            await interaction.response.send_message("invalid link")
        else:
            data = modules.karuta_tag_sorter.sorter.karuta_series(link=link)
            with open(f"{interaction.user.id}.txt", "w", encoding="utf-8") as file:
                file.write(data)
            await interaction.response.send_message(file=nextcord.File(fp=f"{interaction.user.id}.txt"))
            os.remove(f"{interaction.user.id}.txt")

            
def setup(bot: commands.Bot):
    bot.add_cog(Karuta(bot))